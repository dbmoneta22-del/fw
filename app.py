
from flask import Flask, request, send_file, render_template_string, redirect, url_for
from docx import Document
from docx.shared import Pt
from PIL import Image
import pytesseract, uuid, os
from datetime import datetime

app = Flask(__name__)
REPORTS_DIR = "reports"
os.makedirs(REPORTS_DIR, exist_ok=True)

def split_text(text):
    parts = {
        "lavoro": [],
        "telefoni": [],
        "immobili": [],
        "veicoli": [],
        "resto": []
    }
    for line in text.splitlines():
        l = line.lower()
        if any(k in l for k in ["lavora", "azienda", "impiego"]):
            parts["lavoro"].append(line)
        elif any(k in l for k in ["tel", "cell", "+39"]):
            parts["telefoni"].append(line)
        elif any(k in l for k in ["via", "residente", "immobile"]):
            parts["immobili"].append(line)
        elif any(k in l for k in ["auto", "targa", "veicolo"]):
            parts["veicoli"].append(line)
        else:
            parts["resto"].append(line)
    return parts

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST" and "reset" in request.form:
        return redirect(url_for("index"))

    data = {
        "nome":"", "cf":"", "nascita":"", "luogo":"", "residenza":"",
        "intervista":"", "lavoro":"", "telefoni":"", "immobili":"", "veicoli":""
    }

    highlighted = False

    if request.method == "POST":
        for k in data:
            data[k] = request.form.get(k,"")

        file = request.files.get("image")
        if file and file.filename:
            img = Image.open(file.stream).convert("RGB")
            img.thumbnail((2000,2000))
            testo_ocr = pytesseract.image_to_string(img, lang="ita", config="--psm 6")
            if not data["intervista"]:
                data["intervista"] = testo_ocr
                highlighted = True

        if "smista" in request.form:
            parts = split_text(data["intervista"])
            data["lavoro"] = "\n".join(parts["lavoro"])
            data["telefoni"] = "\n".join(parts["telefoni"])
            data["immobili"] = "\n".join(parts["immobili"])
            data["veicoli"] = "\n".join(parts["veicoli"])
            data["intervista"] = "\n".join(parts["resto"])
            highlighted = False
            return render_template_string(FORM_HTML, data=data, highlighted=highlighted)

        if "preview" in request.form:
            return render_template_string(PREVIEW_HTML, data=data)

        version = len(os.listdir(REPORTS_DIR)) + 1

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        doc.add_heading("RAPPORTO INFORMATIVO OSINT", 0)
        doc.add_paragraph(f"Versione: {version}")
        doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

        def section(title, keys):
            doc.add_heading(title, level=1)
            t = doc.add_table(rows=1, cols=2)
            t.style = "Table Grid"
            for k in keys:
                r = t.add_row().cells
                r[0].text = k.replace('_',' ').upper()
                r[1].text = data[k] or "—"

        section("DATI ANAGRAFICI", ["nome","cf","nascita","luogo","residenza"])
        section("DATI INVESTIGATI OSINT", ["intervista","lavoro","telefoni","immobili","veicoli"])

        path = os.path.join(REPORTS_DIR, f"report_v{version}_{uuid.uuid4().hex[:6]}.docx")
        doc.save(path)
        return send_file(path, as_attachment=True)

    return render_template_string(FORM_HTML, data=data, highlighted=highlighted)

FORM_HTML = """
<!doctype html>
<html lang="it">
<head>
<meta charset="utf-8">
<title>Generatore Rapporto OSINT</title>
<style>
body{background:#f1f5f9;color:#0f172a;font-family:system-ui}
.wrapper{max-width:1100px;margin:40px auto;background:white;padding:40px;border-radius:20px;box-shadow:0 20px 40px rgba(0,0,0,.15)}
h1{color:#0369a1}
.section{margin-top:30px}
.section-title{margin-bottom:10px;font-weight:700;color:#334155}
.grid{display:grid;grid-template-columns:repeat(2,1fr);gap:14px}
input, textarea{padding:12px;border-radius:10px;border:1px solid #cbd5f5;background:white;color:#0f172a}
textarea{grid-column:1/3;min-height:100px}
.highlight{background:#fef08a}
.actions{display:flex;gap:14px;margin-top:30px}
button{flex:1;padding:16px;border-radius:12px;border:none;font-weight:700;font-size:15px;cursor:pointer}
.reset{background:#e5e7eb;color:#0f172a}
.smista{background:#f59e0b;color:white}
.preview{background:#38bdf8;color:white}
.generate{background:#22c55e;color:white}
</style>
</head>
<body>
<div class="wrapper">
<h1>Generatore Rapporto OSINT</h1>
<form method="POST" enctype="multipart/form-data">

<div class="section">
<div class="section-title">DATI ANAGRAFICI (Open Source Intelligence)</div>
<div class="grid">
<input name="nome" placeholder="Nome e Cognome" value="{{data.nome}}">
<input name="cf" placeholder="Codice Fiscale" value="{{data.cf}}">
<input name="nascita" placeholder="Data di nascita" value="{{data.nascita}}">
<input name="luogo" placeholder="Luogo di nascita" value="{{data.luogo}}">
<input name="residenza" placeholder="Residenza" value="{{data.residenza}}">
</div>
</div>

<div class="section">
<div class="section-title">DATI INVESTIGATI OSINT</div>
<div class="grid">
<textarea name="intervista" class="{{'highlight' if highlighted else ''}}" placeholder="Intervista Web (OCR)">{{data.intervista}}</textarea>
<textarea name="lavoro" placeholder="Rapporto lavorativo">{{data.lavoro}}</textarea>
<textarea name="telefoni" placeholder="Numeri di telefono">{{data.telefoni}}</textarea>
<textarea name="immobili" placeholder="Immobili intestati">{{data.immobili}}</textarea>
<textarea name="veicoli" placeholder="Veicoli intestati">{{data.veicoli}}</textarea>
</div>
</div>

<div class="section">
<div class="section-title">DOCUMENTO</div>
<input type="file" name="image" accept="image/*">
</div>

<div class="actions">
<button name="reset" class="reset">Reset</button>
<button name="smista" class="smista">Smista testo</button>
<button name="preview" class="preview">Anteprima</button>
<button type="submit" class="generate">Genera Report</button>
</div>

</form>
</div>
</body>
</html>
"""

PREVIEW_HTML = """
<!doctype html>
<html lang="it">
<head>
<meta charset="utf-8">
<title>Anteprima</title>
<style>
body{background:#f1f5f9;color:#0f172a;font-family:system-ui}
table{width:90%;margin:40px auto;border-collapse:collapse;background:white}
td{border:1px solid #cbd5f5;padding:12px}
</style>
</head>
<body>
<h2 style="text-align:center">Anteprima Rapporto OSINT</h2>
<table>
{% for k,v in data.items() %}
<tr>
<td>{{k.replace('_',' ').upper()}}</td>
<td>{{v or '—'}}</td>
</tr>
{% endfor %}
</table>
<form method="POST" style="text-align:center">
{% for k,v in data.items() %}
<input type="hidden" name="{{k}}" value="{{v}}">
{% endfor %}
<button type="submit" class="generate">Conferma e genera report</button>
</form>
</body>
</html>
"""
